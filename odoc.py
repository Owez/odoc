from pathlib import Path
from docx import Document
import re
from docx.shared import Inches


class ODoc:
    _level_cache = 0

    def __init__(self, md: str, verbose: bool = False) -> None:
        """Creates a new ODoc document from inputted markdown content"""
        # add basics
        self.lines = md.splitlines()
        self.doc = Document()
        self.verbose = verbose

        # parse
        self._parse()

    def save(self, path: Path):
        """Saves ODoc document to the provided `path` filepath"""
        self.doc.save(path)

    def _info(self, msg: str, tabbing: str = "  "):
        """Logs infomation to the terminal if `self.verbose` is set"""

        if self.verbose:
            print(f"{tabbing}{msg}..")

    def _parse(self):
        """Parses provided markdown content inside of `self.lines` into in-memory docx file"""

        self._info("Generating document", "")

        # line parsing
        skip = 0
        for ind, line in enumerate(self.lines):
            # skipping
            if skip != 0:
                skip -= 1
                continue

            # matching
            if len(line) == 0:
                # empty line, skip
                continue
            elif line.startswith("#"):
                # heading
                self._add_heading(line)
            elif line.startswith("-"):
                # parse bulletpoint indent block
                skip = self._parse_indents(ind, True)
            elif _str_is_listpoint(line):
                # parse list item indent block
                skip = self._parse_indents(ind, False)
            else:
                # paragraph
                self._add_paragraph(line)

        # add header numbering
        _add_heading_numbering(self.doc)

    def _parse_indents(
        self,
        from_ind: int,
        bulletpoint: bool,
    ) -> int:
        """Parses indentation blocks starting with a bulletpoint if not a list item, returning how many lines to skip"""

        def strip_listing(line: str, bulletpoint: bool) -> str:
            """Removes list item or bulletpoint"""

            line = line.strip()
            line = line[1:] if bulletpoint else line[2:]
            return line.strip()

        (level, self._level_cache) = _get_level(self.lines[from_ind], self._level_cache)
        first_line = strip_listing(self.lines[from_ind], bulletpoint)
        lines = [first_line]

        for further_ind, further_line in enumerate(self.lines[from_ind + 1 :]):
            if further_line.startswith(" "):
                # indented
                trimmed = further_line.strip()
                new_bulletpoint = trimmed.startswith("-")
                if new_bulletpoint or _str_is_listpoint(trimmed):
                    # new bulletpoint or listpoint
                    self._parse_indents_finish(lines, level, bulletpoint)

                    skipped = len(lines)
                    skipped += self._parse_indents(
                        from_ind + further_ind + 1, new_bulletpoint
                    )

                    return skipped
                else:
                    # new paragraph
                    lines.append(further_line)
            else:
                # not indented
                break

        self._parse_indents_finish(lines, level, bulletpoint)
        return len(lines) - 1

    def _parse_indents_finish(self, lines: list, level: int, bulletpoint: bool):
        """Finishes up indent parsing from `self._parse_indents` by writing to docx"""

        if bulletpoint:
            self._add_bulletpoint(lines, level)
        else:
            self._add_listpoint(lines, level)

    def _add_heading(self, line: str):
        """Adds a heading from raw markdown to the document"""

        level = _count_chars(line, "#")
        text = line[level:].strip()

        self._info(f"Adding heading (" + "#" * (level + 1) + ")")
        self.doc.add_heading(text, level)

    def _add_bulletpoint(self, lines: list, level: int):
        """Adds a bulletpoint for provided lines at a given level"""

        self._info(f"Adding level {level} bulletpoint")

        level_style = "" if level == 0 else f" {level + 1}"
        paragraph = self.doc.add_paragraph("", style="List Bullet" + level_style)
        self._add_paragraph_runs(lines[0], paragraph)  # TODO: multiple lines

    def _add_listpoint(self, lines: list, level: int):
        """Adds a listpoint (ordered list) for provided lines at a given level"""

        self._info(f"Adding level {level} listpoint")

        level_style = "" if level == 0 else f" {level + 1}"
        paragraph = self.doc.add_paragraph("", style="List Number" + level_style)
        self._add_paragraph_runs(lines[0], paragraph)  # TODO: multiple lines

    def _add_paragraph(self, line: str):
        """Adds one plain (pure) paragraph of text with styling only from inline markdown"""

        self._info("Adding paragraph")

        paragraph = self.doc.add_paragraph()
        self._add_paragraph_runs(line, paragraph)

    def _add_paragraph_runs(self, line: str, paragraph):
        """Adds paragraph runs to an existing docx paragraph for a line"""

        self._info("Adding next paragraph run", "    ")

        paragraph.add_run(line.strip())  # TODO: parse line, make sure to strip
        pass  # TODO: parse paragraph bits


def _iter_heading(paragraphs):
    for paragraph in paragraphs:
        isItHeading = re.match("Heading ([1-9])", paragraph.style.name)
        if isItHeading:
            yield int(isItHeading.groups()[0]), paragraph


def _add_heading_numbering(document):
    hNums = [0, 0, 0, 0, 0]
    for index, hx in _iter_heading(document.paragraphs):
        # ---put zeroes below---
        for i in range(index + 1, 5):
            hNums[i] = 0
        # ---increment this---
        hNums[index] += 1
        # ---prepare the string---
        hStr = ""
        for i in range(1, index + 1):
            hStr += "%d." % hNums[i]
        # ---add the numbering---
        hx.text = hStr + " " + hx.text


def _count_chars(line: str, target: str) -> int:
    """Counts number of prefixed characters of `target` for a provided `line` string"""

    count = 0
    for c in line:
        if c == target:
            count += 1
        else:
            break
    return count


def _get_level(line: str, level_cache: int = 0) -> tuple:
    """Gets level with the context of the first amount of spaces, returning the formatted level and raw level cache"""

    level = _count_chars(line, " ")

    if level_cache == 0 and level > 1:
        level_cache = level
        level = 1
    elif level != 0:
        level = level / level_cache

    return (level, level_cache)


def _str_is_listpoint(string: str) -> bool:
    """Checks if a provided string (line) is a listpoint"""

    return len(string) > 1 and string[0].isnumeric() and string[1] == "."
